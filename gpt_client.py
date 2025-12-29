import os
import json
import time
import requests
import numpy as np
import faiss
import pdfplumber
from docx import Document
import pytesseract
import cv2

from certificate_ocr import extract_certificates_from_notes

# ================= CONFIG =================
OLLAMA_API = "http://localhost:11434/api"
CHAT_MODEL = "llama3"
EMBED_MODEL = "nomic-embed-text"

NOTES_PATH = r"C:\Users\chris\OneDrive\NOTES"
CACHE_DIR = "rag_cache"
INDEX_PATH = os.path.join(CACHE_DIR, "index.faiss")
CHUNKS_PATH = os.path.join(CACHE_DIR, "chunks.json")

os.makedirs(CACHE_DIR, exist_ok=True)

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ================= PERFORMANCE TUNING =================
FAISS_STRONG = 350
FAISS_WEAK = 450
TOP_K_CHUNKS = 2           # 🔥 reduce context
MAX_TOKENS = 200           # 🔥 limit LLM output

GK_KEYWORDS = [
    "what is", "explain", "define", "history",
    "advantages", "disadvantages", "types of"
]

# ================= GLOBAL =================
LAST_FAISS_DISTANCE = None
EMBED_CACHE = {}           # 🔥 embedding cache

# ================= FILE READERS =================
def read_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_pdf(path):
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    except:
        pass
    return text

def read_docx(path):
    text = []
    try:
        doc = Document(path)
        for p in doc.paragraphs:
            if p.text.strip():
                text.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
    except:
        pass
    return "\n".join(text)

# ================= OCR (NOTES) =================
def preprocess_image(img_path):
    img = cv2.imread(img_path)
    if img is None:
        return None
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.resize(gray, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
    _, gray = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
    return gray

def read_images(path):
    text = ""
    for root, _, files in os.walk(path):
        for file in files:
            if file.lower().endswith((".png", ".jpg", ".jpeg")):
                img_path = os.path.join(root, file)
                img = preprocess_image(img_path)
                if img is not None:
                    t = pytesseract.image_to_string(img, config="--psm 6")
                    if t.strip():
                        text += t + "\n"
    return text

def load_documents():
    docs = []
    for root, _, files in os.walk(NOTES_PATH):
        for file in files:
            full = os.path.join(root, file)
            try:
                if file.lower().endswith(".txt"):
                    docs.append(read_txt(full))
                elif file.lower().endswith(".pdf"):
                    docs.append(read_pdf(full))
                elif file.lower().endswith(".docx"):
                    docs.append(read_docx(full))
            except:
                pass

    docs.append(read_images(NOTES_PATH))
    return docs

# ================= CHUNKING =================
def chunk_text(text, size=300, overlap=80):
    chunks = []
    start = 0
    while start < len(text):
        chunks.append(text[start:start + size])
        start += size - overlap
    return chunks

# ================= EMBEDDINGS (CACHED) =================
def get_embedding(text):
    if text in EMBED_CACHE:
        return EMBED_CACHE[text]

    r = requests.post(
        f"{OLLAMA_API}/embeddings",
        json={"model": EMBED_MODEL, "prompt": text}
    )
    emb = np.array(r.json()["embedding"], dtype="float32")
    EMBED_CACHE[text] = emb
    return emb

# ================= VECTOR STORE =================
def build_vector_store():
    print("🔄 Building knowledge base from NOTES...")
    documents = load_documents()
    chunks = []

    for d in documents:
        chunks.extend(chunk_text(d))

    embeddings = [get_embedding(c) for c in chunks]
    dim = len(embeddings[0])

    index = faiss.IndexFlatL2(dim)
    index.add(np.array(embeddings))

    faiss.write_index(index, INDEX_PATH)
    with open(CHUNKS_PATH, "w", encoding="utf-8") as f:
        json.dump(chunks, f)

    print(f"✅ Indexed {len(chunks)} chunks")
    return index, chunks

def load_vector_store():
    if os.path.exists(INDEX_PATH) and os.path.exists(CHUNKS_PATH):
        return (
            faiss.read_index(INDEX_PATH),
            json.load(open(CHUNKS_PATH, "r", encoding="utf-8"))
        )
    return build_vector_store()

# ================= RETRIEVAL =================
def retrieve_context(query, index, chunks):
    global LAST_FAISS_DISTANCE

    q_emb = get_embedding(query)
    D, I = index.search(np.array([q_emb]), 5)

    LAST_FAISS_DISTANCE = float(D[0][0])
    print("🔍 FAISS distances:", D[0])

    return "\n".join(chunks[i] for i in I[0][:TOP_K_CHUNKS])

# ================= STARTUP =================
startup_start = time.time()

index, chunks = load_vector_store()
certificates = extract_certificates_from_notes(NOTES_PATH)

print(f"🎓 Certificates detected: {len(certificates)}")
print(f"⏱ Startup completed in {time.time() - startup_start:.2f}s")

print("\n🧠 Mini GPT Chat (FAST MODE)")
print("Type 'exit' to quit.\n")

# ================= CHAT LOOP =================
while True:
    user_input = input("You: ").strip()
    if user_input.lower() == "exit":
        break

    start_q = time.time()
    user_lower = user_input.lower()

    # ===== CERTIFICATE LIST =====
    if user_lower in [
        "do i have a certificate",
        "do i have certificates",
        "list my certificates",
        "what certificates are in my onedrive"
    ]:
        if certificates:
            print("AI: ✅ Certificates found:")
            for c in certificates:
                status = "VALID" if c["verification"]["is_valid"] else "INVALID"
                print(f"  📄 {c['file']} → {status}")
        else:
            print("AI: ❌ No certificates found")

        print(f"\n⏱ Answer time: {time.time() - start_q:.2f}s\n")
        continue

    # ===== SMART GK SKIP =====
    if any(user_lower.startswith(k) for k in GK_KEYWORDS):
        context = ""
        LAST_FAISS_DISTANCE = None
    else:
        context = retrieve_context(user_input, index, chunks)

    # ===== DECISION =====
    if context.strip() and LAST_FAISS_DISTANCE is not None:
        if LAST_FAISS_DISTANCE < FAISS_STRONG:
            use_notes = True
        elif LAST_FAISS_DISTANCE < FAISS_WEAK:
            use_notes = True
        else:
            use_notes = False
    else:
        use_notes = False

    # ===== PROMPT =====
    if use_notes:
        prompt = f"""
Answer using ONLY the OneDrive NOTES below.
If not explicitly present, say so.

NOTES:
{context}

QUESTION:
{user_input}
"""
    else:
        prompt = f"""
The answer is NOT present in OneDrive NOTES.
First say exactly: ❌ Not in OneDrive
Then answer using general knowledge.

QUESTION:
{user_input}
"""

    # ===== LLM CALL (LIMITED) =====
    llm_start = time.time()
    r = requests.post(
        f"{OLLAMA_API}/generate",
        json={
            "model": CHAT_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {
                "num_predict": MAX_TOKENS,
                "temperature": 0.2
            }
        }
    )

    print("AI:", r.json()["response"])
    print(f"\n⏱ Total answer time: {time.time() - start_q:.2f}s\n")
